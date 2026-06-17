import streamlit as st
import zipfile
import os
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile

from staticmap import StaticMap, CircleMarker


st.title("📍 Konwerter KMZ → Word (raport drzew)")


# =========================
# MAPA
# =========================
def get_map(lat, lon, i):
    try:
        m = StaticMap(800, 500)
        marker = CircleMarker((lon, lat), 'red', 12)
        m.add_marker(marker)

        image = m.render(zoom=17)

        path = f"map_{i}.png"
        image.save(path)

        return path
    except Exception as e:
        st.error(f"Błąd mapy: {e}")
        return None


# =========================
# ZDJĘCIA
# =========================
def get_images_for_placemark(soup, extract_dir):
    images = []
    for img in soup.find_all("img"):
        src = img.get("src")
        if src and "cloud_media" in src:
            filename = src.split("/")[-1]
            for r, d, f in os.walk(extract_dir):
                if filename in f:
                    images.append(os.path.join(r, filename))
                    break
    return list(dict.fromkeys(images))


# =========================
# UPLOAD
# =========================
uploaded_file = st.file_uploader("Wgraj plik KMZ", type=["kmz"])


if uploaded_file is not None:

    if st.button("🚀 Generuj Word"):

        with st.spinner("Przetwarzanie..."):

            with tempfile.TemporaryDirectory() as tmpdir:

                kmz_path = os.path.join(tmpdir, "file.kmz")

                with open(kmz_path, "wb") as f:
                    f.write(uploaded_file.read())

                extract_dir = os.path.join(tmpdir, "kmz")
                os.makedirs(extract_dir, exist_ok=True)

                # UNZIP
                with zipfile.ZipFile(kmz_path, 'r') as z:
                    z.extractall(extract_dir)

                # FIND KML
                kml_file = None
                for r, d, f in os.walk(extract_dir):
                    for file in f:
                        if file.endswith(".kml"):
                            kml_file = os.path.join(r, file)

                if not kml_file:
                    st.error("Nie znaleziono KML")
                    st.stop()

                # PARSE
                ns = {'kml': 'http://www.opengis.net/kml/2.2'}
                tree = ET.parse(kml_file)
                root = tree.getroot()
                placemarks = root.findall('.//kml:Placemark', ns)

                # WORD
                doc = Document()
                doc.add_heading("RAPORT OBSZAROWY DRZEW", 1)

                # LOOP
                for i, pm in enumerate(placemarks, 1):

                    name = pm.find('kml:name', ns)
                    name = name.text if name is not None else f"Drzewo {i}"
                    doc.add_heading(name, 2)

                    coords = pm.find('.//kml:coordinates', ns)

                    if coords is not None:
                        lon, lat = coords.text.strip().split(",")[:2]
                        lat, lon = float(lat), float(lon)

                        doc.add_paragraph(f"GPS: {lat}, {lon}")

                        map_img = get_map(lat, lon, i)
                        if map_img:
                            doc.add_picture(map_img, width=Inches(5))

                    desc = pm.find('kml:description', ns)

                    if desc is not None:
                        soup = BeautifulSoup(desc.text, "html.parser")

                        pre = soup.find("pre")
                        if pre:
                            doc.add_paragraph("Opis: " + pre.text)

                        table = soup.find("table")
                        if table:
                            t = doc.add_table(rows=1, cols=2)
                            t.style = "Table Grid"
                            t.rows[0].cells[0].text = "Parametr"
                            t.rows[0].cells[1].text = "Wartość"

                            for row in table.find_all("tr"):
                                cols = row.find_all("td")
                                if len(cols) == 2:
                                    r = t.add_row().cells
                                    r[0].text = cols[0].text.strip()
                                    r[1].text = cols[1].text.strip()

                        doc.add_paragraph("Zdjęcia:")

                        images = get_images_for_placemark(soup, extract_dir)

                        if images:
                            for img_path in images:
                                try:
                                    doc.add_picture(img_path, width=Inches(4))
                                except:
                                    doc.add_paragraph("Błąd zdjęcia")
                        else:
                            doc.add_paragraph("Brak zdjęć")

                    doc.add_page_break()

                # SAVE
                output_path = os.path.join(tmpdir, "raport.docx")
                doc.save(output_path)

                # DOWNLOAD
                with open(output_path, "rb") as f:
                    st.download_button(
                        "⬇️ Pobierz Word",
                        f,
                        file_name="raport_drzew.docx"
                    )

        st.success("Gotowe!")